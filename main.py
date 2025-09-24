from fastapi import FastAPI, Form, HTTPException, Request, Depends
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel
from datetime import datetime
import docx
import os
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from bson import ObjectId
import dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from fastapi.security import OAuth2PasswordBearer
from passlib.context import CryptContext
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
dotenv.load_dotenv()
MONGO_URI = os.getenv("MONGO_URI")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

app = FastAPI()

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")

# Templates
templates = Jinja2Templates(directory="templates")

# MongoDB connection
client = AsyncIOMotorClient(MONGO_URI)
db = client["mom_platform"]
users_collection = db["users"]
reports_collection = db["reports"]
action_items_collection = db["action_items"]
admins_collection = db["admins"]

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="login")

# Pydantic models
class Report(BaseModel):
    date: str
    name: str
    yesterday: str
    today: str
    blockers: str

class User(BaseModel):
    name: str

class Admin(BaseModel):
    email: str
    password: str

class ActionItems(BaseModel):
    date: str
    items: str

# Helper function to get today's date
def get_today_date():
    return datetime.now().strftime("%Y-%m-%d")

# Helper function to verify password
def verify_password(plain_password, hashed_password):
    try:
        return pwd_context.verify(plain_password, hashed_password)
    except Exception as e:
        logger.error(f"Password verification failed: {e}")
        return False

# Helper function to authenticate admin
async def authenticate_admin(email: str, password: str):
    logger.info(f"Attempting to authenticate admin with email: {email}")
    admin = await admins_collection.find_one({"email": email})
    if not admin:
        logger.warning(f"No admin found for email: {email}")
        return False
    if not verify_password(password, admin["password"]):
        logger.warning(f"Invalid password for email: {email}")
        return False
    logger.info(f"Authentication successful for email: {email}")
    return True

# Get current admin from cookies
async def get_current_admin(request: Request):
    token = request.cookies.get("access_token")
    if not token:
        raise HTTPException(
            status_code=401,
            detail="Not authenticated",
            headers={"WWW-Authenticate": "Bearer", "Location": "/login"}
        )
    admin = await admins_collection.find_one({"email": token})
    if not admin:
        raise HTTPException(
            status_code=401,
            detail="Invalid authentication credentials",
            headers={"WWW-Authenticate": "Bearer", "Location": "/login"}
        )
    return token

# Login endpoint
@app.post("/login")
async def login(email: str = Form(...), password: str = Form(...)):
    logger.info(f"Login attempt for email: {email}")
    if await authenticate_admin(email, password):
        access_token = email
        resp = JSONResponse({"access_token": access_token, "token_type": "bearer"})
        resp.set_cookie(key="access_token", value=access_token, httponly=True, samesite="lax", max_age=86400)
        return resp
    raise HTTPException(status_code=401, detail="Invalid email or password")

# Login page
@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    return templates.TemplateResponse("login.html", {"request": request})

# Home page
@app.get("/", response_class=HTMLResponse)
async def home_page(request: Request):
    users = await users_collection.find().to_list(None)
    user_names = [user["name"] for user in users]
    return templates.TemplateResponse("index.html", {
        "request": request,
        "today_date": get_today_date(),
        "users": user_names
    })

# Check if report exists
@app.get("/check_report/{date}/{name}")
async def check_report(date: str, name: str):
    report = await reports_collection.find_one({"date": date, "name": name})
    if report:
        return {
            "exists": True,
            "report": {
                "yesterday": report["yesterday"],
                "today": report["today"],
                "blockers": report["blockers"],
                "id": str(report["_id"])
            }
        }
    return {"exists": False}

# Save report
@app.post("/save_report")
async def save_report(
    date: str = Form(...),
    name: str = Form(...),
    yesterday: str = Form(...),
    today: str = Form(...),
    blockers: str = Form(...)
):
    # Word count validation (backend fallback)
    def count_words(text):
        return len(text.strip().split())
    
    if count_words(yesterday) > 10 or count_words(today) > 10:
        raise HTTPException(status_code=400, detail="Each field must not exceed 10 words")
    
    if not yesterday.strip() or not today.strip():
        raise HTTPException(status_code=400, detail="Yesterday's tasks and today's priorities cannot be empty")
    
    existing_report = await reports_collection.find_one({"date": date, "name": name})
    if existing_report:
        raise HTTPException(status_code=400, detail="Report already exists for this user and date")
    
    report = Report(date=date, name=name, yesterday=yesterday, today=today, blockers=blockers)
    await reports_collection.insert_one(report.dict())
    return {"message": "Report saved successfully"}

# Admin page (redirect to team members)
@app.get("/admin", response_class=HTMLResponse)
async def admin_page(request: Request, token: str = Depends(get_current_admin)):
    return templates.TemplateResponse("team_members.html", {"request": request, "today_date": get_today_date()})

# Team Members page
@app.get("/team_members", response_class=HTMLResponse)
async def team_members_page(request: Request, token: str = Depends(get_current_admin)):
    return templates.TemplateResponse("team_members.html", {"request": request, "today_date": get_today_date()})

# Reports page
@app.get("/reports", response_class=HTMLResponse)
async def reports_page(request: Request, token: str = Depends(get_current_admin)):
    return templates.TemplateResponse("reports.html", {"request": request, "today_date": get_today_date()})

# Add user
@app.post("/add_user")
async def add_user(name: str = Form(...), token: str = Depends(get_current_admin)):
    existing_user = await users_collection.find_one({"name": name})
    if existing_user:
        raise HTTPException(status_code=400, detail="User already exists")
    user = User(name=name)
    await users_collection.insert_one(user.dict())
    return {"message": "User added successfully"}

# Delete user
@app.post("/delete_user")
async def delete_user(name: str = Form(...), token: str = Depends(get_current_admin)):
    existing_user = await users_collection.find_one({"name": name})
    if not existing_user:
        raise HTTPException(status_code=404, detail="User not found")
    await users_collection.delete_one({"name": name})
    # Optionally, delete associated reports
    await reports_collection.delete_many({"name": name})
    return {"message": "User deleted successfully"}

# Get users
@app.get("/users")
async def get_users(token: str = Depends(get_current_admin)):
    users = await users_collection.find().to_list(None)
    return [user["name"] for user in users]

# Get reports by date
@app.get("/reports/{date}")
async def get_reports(date: str, token: str = Depends(get_current_admin)):
    reports = await reports_collection.find({"date": date}).to_list(None)
    return [
        {
            "sno": i + 1,
            "name": report["name"],
            "yesterday": report["yesterday"],
            "today": report["today"],
            "blockers": report["blockers"]
        } for i, report in enumerate(reports)
    ]

# Generate and save action items (bullet format)
@app.get("/generate_action_items/{date}")
async def generate_action_items(date: str, token: str = Depends(get_current_admin)):
    reports = await reports_collection.find({"date": date}).to_list(None)
    if not reports:
        raise HTTPException(status_code=404, detail="No reports found for this date")
    
    existing_action_items = await action_items_collection.find_one({"date": date})
    if existing_action_items:
        return {"action_items": existing_action_items["items"]}
    
    content = "Daily Reports:\n"
    for report in reports:
        content += f"{report['name']}:\n- Yesterday: {report['yesterday']}\n- Today: {report['today']}\n- Blockers: {report['blockers']}\n\n"
    
    prompt = f"""From the following daily reports, extract the main and unique action items based on the 'Today's Priorities' field. 
    Follow these guidelines:
    - Extract each distinct task as a separate action item - do not combine unrelated tasks, even if they come from the same person.
    - For example, if one person mentions both "lead generation" and "MOM automation implementation", treat them as two separate action items.
    - Only consolidate tasks that are directly related (e.g., multiple aspects of the same feature like "enhancing and testing" the same agents).
    - Focus on the primary objectives for the day.
    - Each action item should be distinct and represent a unique goal or task.
    - Output in bullet point format using this structure:
    • [Action Item]
    - Ensure no unrelated tasks are merged together.

    Reports:\n{content}"""
    
    try:
        llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", google_api_key=GEMINI_API_KEY)
        response = await llm.ainvoke(prompt)
        
        action_items_text = response.content.strip()
        action_items = ActionItems(date=date, items=action_items_text)
        await action_items_collection.insert_one(action_items.dict())
        
        return {"action_items": action_items_text}
    except Exception as e:
        logger.error(f"Error generating action items: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate action items")

# Get action items
@app.get("/action_items/{date}")
async def get_action_items(date: str, token: str = Depends(get_current_admin)):
    action_items = await action_items_collection.find_one({"date": date})
    if action_items:
        return {"action_items": action_items["items"]}
    return {"action_items": ""}

# Download consolidated report
@app.get("/download_report/{date}")
async def download_report(date: str, token: str = Depends(get_current_admin)):
    reports = await reports_collection.find({"date": date}).to_list(None)
    action_items = await action_items_collection.find_one({"date": date})
    if not reports:
        raise HTTPException(status_code=404, detail="No reports found for this date")
    
    doc = docx.Document()
    doc.add_heading(f"MOM Consolidated Report - {date}", 0)
    
    # Reports section
    doc.add_heading("Daily Progress Reports", level=1)
    table = doc.add_table(rows=len(reports) + 1, cols=5)
    table.style = "Table Grid"
    
    headers = ["S.No", "Team Member", "Yesterday's Tasks", "Today's Priorities", "Blockers"]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    
    for i, report in enumerate(reports, 1):
        row = table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = report["name"]
        row.cells[2].text = report["yesterday"]
        row.cells[3].text = report["today"]
        row.cells[4].text = report["blockers"] or "None"
    
    # Action items section
    if action_items and action_items["items"]:
        doc.add_heading("Action Items", level=1)
        action_paragraph = doc.add_paragraph()
        action_paragraph.add_run(action_items["items"])
    
    file_path = f"mom-report-{date}.docx"
    doc.save(file_path)
    return FileResponse(file_path, filename=file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Initialize admin
@app.on_event("startup")
async def startup_event():
    admin_email = "aravind@gmail.com"
    admin_password = "Admin@1234"
    admin = await admins_collection.find_one({"email": admin_email})
    if not admin:
        hashed_password = pwd_context.hash(admin_password)
        await admins_collection.insert_one({"email": admin_email, "password": hashed_password})
        logger.info(f"Default admin created: {admin_email} / {admin_password}")